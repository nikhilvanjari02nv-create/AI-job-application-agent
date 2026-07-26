# AI Resume Generator
# Placeholder: full implementation scaffold
from pathlib import Path
from docx import Document
from app.models.job import Job
from app.agents.resume_parser import parse_resume
from app.services.llm import ask_llm

BASE_DIR = Path(__file__).resolve().parent.parent
MASTER_RESUME = BASE_DIR / "templates" / "master_resume.docx"
PROMPT_FILE = BASE_DIR / "prompts" / "resume_tailor.txt"
OUTPUT_DIR = BASE_DIR / "generated_resumes"
OUTPUT_DIR.mkdir(exist_ok=True)

def generate_resume(job: Job):
    if not MASTER_RESUME.exists():
        raise FileNotFoundError(MASTER_RESUME)

    structured = parse_resume(str(MASTER_RESUME))
    prompt = PROMPT_FILE.read_text(encoding="utf-8")

    llm_prompt = f"""{prompt}

JOB TITLE:
{job.title}

COMPANY:
{job.company}

JOB DESCRIPTION:
{job.description or ""}

RESUME:
{structured}
"""

    tailored = ask_llm(llm_prompt)

    doc = Document()
    for line in tailored.splitlines():
        line=line.strip()
        if not line:
            continue
        if line.startswith("#"):
            doc.add_heading(line.lstrip("# ").strip(), level=1)
        elif line.startswith("-"):
            doc.add_paragraph(line[1:].strip(), style="List Bullet")
        else:
            doc.add_paragraph(line)

    filename=f"{job.company}_{job.title}".replace("/","_").replace("\\","_").replace(" ","_")
    out=OUTPUT_DIR/f"{filename}.docx"
    doc.save(out)
    return str(out)
